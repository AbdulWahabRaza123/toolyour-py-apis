"""
Document conversion service.
Handles DOCX to PDF and other document conversions.
"""

import io
import os
import tempfile
from typing import Optional
from docx import Document
from reportlab.lib.pagesizes import A4, letter, legal
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import structlog

from .types import ServiceResponse, FileInput, ConversionOptions

logger = structlog.get_logger(__name__)


class DocumentConverterService:
    """Service for converting documents between formats."""

    def __init__(self):
        self.supported_conversions = {
            'docx': ['pdf', 'txt', 'html'],
            'pdf': ['txt'],
            'txt': ['pdf', 'docx'],
        }

    def can_convert(self, source_format: str, target_format: str) -> bool:
        """Check if conversion is supported."""
        source_format = source_format.lower().replace('.', '')
        target_format = target_format.lower().replace('.', '')
        
        return target_format in self.supported_conversions.get(source_format, [])

    def get_supported_formats(self, source_format: str) -> list:
        """Get supported target formats for a source format."""
        source_format = source_format.lower().replace('.', '')
        return self.supported_conversions.get(source_format, [])

    async def convert_docx_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """
        Convert DOCX file to PDF.
        
        Args:
            file_buffer: The DOCX file content as bytes
            options: Conversion options (page size, orientation, etc.)
            
        Returns:
            ServiceResponse with PDF file buffer or error
        """
        try:
            if options is None:
                options = ConversionOptions()

            # Load DOCX document
            docx_stream = io.BytesIO(file_buffer)
            doc = Document(docx_stream)

            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            
            # Set page size
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            # Get styles
            styles = getSampleStyleSheet()
            story = []

            # Convert DOCX paragraphs to PDF
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Determine paragraph style based on DOCX style
                    style = self._get_paragraph_style(paragraph, styles)
                    
                    # Add paragraph to story
                    p = Paragraph(self._clean_text(paragraph.text), style)
                    story.append(p)
                    story.append(Spacer(1, 6))

            # Add tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        p = Paragraph(self._clean_text(row_text), styles['Normal'])
                        story.append(p)
                        story.append(Spacer(1, 3))
                story.append(Spacer(1, 12))

            # Build PDF
            pdf_doc.build(story)
            
            # Get PDF content
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info(
                "DOCX to PDF conversion completed",
                output_size=len(pdf_content)
            )

            return ServiceResponse(
                status=200,
                message="DOCX converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("DOCX to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to PDF",
                error=str(e)
            )

    async def convert_docx_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert DOCX file to plain text.
        
        Args:
            file_buffer: The DOCX file content as bytes
            
        Returns:
            ServiceResponse with text content or error
        """
        try:
            # Load DOCX document
            docx_stream = io.BytesIO(file_buffer)
            doc = Document(docx_stream)

            # Extract all text
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text_content.append(row_text)

            full_text = "\n".join(text_content)

            logger.info("DOCX to TXT conversion completed")

            return ServiceResponse(
                status=200,
                message="DOCX converted to TXT successfully",
                data=full_text.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("DOCX to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to TXT",
                error=str(e)
            )

    async def convert_txt_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """
        Convert plain text file to PDF.
        
        Args:
            file_buffer: The text file content as bytes
            options: Conversion options
            
        Returns:
            ServiceResponse with PDF file buffer or error
        """
        try:
            if options is None:
                options = ConversionOptions()

            # Decode text
            text_content = file_buffer.decode('utf-8')

            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            
            # Set page size
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            # Get styles
            styles = getSampleStyleSheet()
            story = []

            # Convert text lines to PDF paragraphs
            for line in text_content.split('\n'):
                if line.strip():
                    p = Paragraph(self._clean_text(line), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))

            # Build PDF
            pdf_doc.build(story)
            
            # Get PDF content
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info("TXT to PDF conversion completed")

            return ServiceResponse(
                status=200,
                message="TXT converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("TXT to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to PDF",
                error=str(e)
            )

    def _get_page_size(self, size_name: str):
        """Get page size tuple based on name."""
        sizes = {
            'A4': A4,
            'letter': letter,
            'legal': legal,
        }
        return sizes.get(size_name, A4)

    def _get_paragraph_style(self, paragraph, styles):
        """Get appropriate PDF style for a DOCX paragraph."""
        # Check if it's a heading
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name.replace('Heading ', '')
            if level == '1':
                return styles['Heading1']
            elif level == '2':
                return styles['Heading2']
            elif level == '3':
                return styles['Heading3']
        
        # Check alignment
        if hasattr(paragraph, 'alignment') and paragraph.alignment:
            if paragraph.alignment == 1:  # Center
                return styles['Normal']
            elif paragraph.alignment == 2:  # Right
                return styles['Normal']
        
        return styles['Normal']

    def _clean_text(self, text: str) -> str:
        """Clean text for PDF generation (escape XML special characters)."""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text


# Global instance
document_converter_service = DocumentConverterService()

