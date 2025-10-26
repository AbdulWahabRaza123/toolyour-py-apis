"""
Document conversion service.
Handles core document formats: DOCX, PDF, TXT, HTML, RTF, ODT, MD.
"""

import io
import os
import tempfile
from typing import Optional
from docx import Document
from reportlab.lib.pagesizes import A4, letter, legal
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import PyPDF2
import structlog

# Additional formats
import markdown
from bs4 import BeautifulSoup
import json

from .types import ServiceResponse, FileInput, ConversionOptions

logger = structlog.get_logger(__name__)


class DocumentConverterService:
    """Service for converting core document formats."""

    def __init__(self):
        self.supported_conversions = {
            # Core document formats only
            'docx': ['pdf', 'txt', 'html', 'rtf', 'odt'],
            'pdf': ['txt', 'docx', 'html', 'rtf'],
            'txt': ['pdf', 'docx', 'html', 'rtf', 'md'],
            'html': ['pdf', 'docx', 'txt', 'md'],
            'rtf': ['pdf', 'docx', 'txt', 'html'],
            'odt': ['pdf', 'docx', 'txt', 'html'],
            'md': ['pdf', 'docx', 'txt', 'html']
        }

    def can_convert(self, source_format: str, target_format: str) -> bool:
        """Check if conversion is supported."""
        source_format = source_format.lower().replace('.', '')
        target_format = target_format.lower().replace('.', '')
        
        return target_format in self.supported_conversions.get(source_format, [])

    def get_supported_formats(self, source_format: str) -> list:
        """Get supported target formats for a source format."""
        source_format = source_format.lower().replace('.', '')
        return self.supported_conversions.get(source_format, [])

    def _get_page_size(self, page_size: str):
        """Get page size for PDF generation."""
        page_sizes = {
            'A4': A4,
            'Letter': letter,
            'Legal': legal
        }
        return page_sizes.get(page_size, A4)

    def _clean_text(self, text: str) -> str:
        """Clean text for PDF generation."""
        # Remove or replace problematic characters
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\r\n', '\n')  # Normalize line endings
        text = text.replace('\r', '\n')  # Normalize line endings
        return text

    # DOCX conversions
    async def convert_docx_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """Convert DOCX to PDF."""
        try:
            if options is None:
                options = ConversionOptions()

            # Load DOCX document
            doc = Document(io.BytesIO(file_buffer))
            
            # Create PDF
            pdf_buffer = io.BytesIO()
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            styles = getSampleStyleSheet()
            story = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Determine style based on paragraph style
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                        style = ParagraphStyle(
                            'CustomHeading',
                            parent=styles['Heading1'],
                            fontSize=18 - (level - 1) * 2,
                            spaceAfter=12,
                        )
                    else:
                        style = styles['Normal']
                    
                    p = Paragraph(self._clean_text(paragraph.text), style)
                    story.append(p)
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info("DOCX to PDF conversion completed")
            return ServiceResponse(
                status=200,
                message="DOCX converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("DOCX to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to PDF",
                error=str(e)
            )

    async def convert_docx_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert DOCX to TXT."""
        try:
            # Load DOCX document
            doc = Document(io.BytesIO(file_buffer))
            
            # Extract text
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            txt_content = '\n'.join(text_content)

            logger.info("DOCX to TXT conversion completed")
            return ServiceResponse(
                status=200,
                message="DOCX converted to TXT successfully",
                data=txt_content.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("DOCX to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to TXT",
                error=str(e)
            )

    async def convert_docx_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert DOCX to HTML."""
        try:
            # Load DOCX document
            doc = Document(io.BytesIO(file_buffer))
            
            # Create HTML
            html_content = ['<html><head><title>Converted Document</title></head><body>']
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                        html_content.append(f'<h{level}>{paragraph.text}</h{level}>')
                    else:
                        html_content.append(f'<p>{paragraph.text}</p>')
            
            html_content.append('</body></html>')
            html_text = '\n'.join(html_content)

            logger.info("DOCX to HTML conversion completed")
            return ServiceResponse(
                status=200,
                message="DOCX converted to HTML successfully",
                data=html_text.encode('utf-8'),
                format="html"
            )

        except Exception as e:
            logger.error("DOCX to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to HTML",
                error=str(e)
            )

    async def convert_docx_to_rtf(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert DOCX to RTF."""
        try:
            # Load DOCX document
            doc = Document(io.BytesIO(file_buffer))
            
            # Create basic RTF
            rtf_content = ['{\\rtf1\\ansi\\deff0']
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                        rtf_content.append(f'{{\\b\\fs{24 - level * 2} {paragraph.text}}}')
                    else:
                        rtf_content.append(paragraph.text)
                    rtf_content.append('\\par')
            
            rtf_content.append('}')
            rtf_text = '\n'.join(rtf_content)

            logger.info("DOCX to RTF conversion completed")
            return ServiceResponse(
                status=200,
                message="DOCX converted to RTF successfully",
                data=rtf_text.encode('utf-8'),
                format="rtf"
            )

        except Exception as e:
            logger.error("DOCX to RTF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to RTF",
                error=str(e)
            )

    async def convert_docx_to_odt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert DOCX to ODT."""
        try:
            # For now, return a placeholder response
            # DOCX to ODT conversion requires additional libraries
            logger.warning("DOCX to ODT conversion requires additional libraries")
            return ServiceResponse(
                status=501,
                message="DOCX to ODT conversion requires additional libraries (odfpy)",
                error="DOCX to ODT conversion not fully implemented"
            )

        except Exception as e:
            logger.error("DOCX to ODT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to ODT",
                error=str(e)
            )

    # PDF conversions
    async def convert_pdf_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert PDF to TXT."""
        try:
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_buffer))
            text_content = []
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            txt_content = '\n'.join(text_content)

            logger.info("PDF to TXT conversion completed")
            return ServiceResponse(
                status=200,
                message="PDF converted to TXT successfully",
                data=txt_content.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("PDF to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting PDF to TXT",
                error=str(e)
            )

    async def convert_pdf_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert PDF to DOCX."""
        try:
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_buffer))
            text_content = []
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            # Create DOCX document
            doc = Document()
            for text in text_content:
                if text.strip():
                    doc.add_paragraph(text)
            
            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()
            docx_buffer.close()

            logger.info("PDF to DOCX conversion completed")
            return ServiceResponse(
                status=200,
                message="PDF converted to DOCX successfully",
                data=docx_content,
                format="docx"
            )

        except Exception as e:
            logger.error("PDF to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting PDF to DOCX",
                error=str(e)
            )

    async def convert_pdf_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert PDF to HTML."""
        try:
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_buffer))
            text_content = []
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            # Create HTML
            html_content = ['<html><head><title>Converted PDF</title></head><body>']
            for text in text_content:
                if text.strip():
                    html_content.append(f'<p>{text}</p>')
            html_content.append('</body></html>')
            
            html_text = '\n'.join(html_content)

            logger.info("PDF to HTML conversion completed")
            return ServiceResponse(
                status=200,
                message="PDF converted to HTML successfully",
                data=html_text.encode('utf-8'),
                format="html"
            )

        except Exception as e:
            logger.error("PDF to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting PDF to HTML",
                error=str(e)
            )

    async def convert_pdf_to_rtf(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert PDF to RTF."""
        try:
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_buffer))
            text_content = []
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            # Create RTF
            rtf_content = ['{\\rtf1\\ansi\\deff0']
            for text in text_content:
                if text.strip():
                    rtf_content.append(text)
                    rtf_content.append('\\par')
            rtf_content.append('}')
            
            rtf_text = '\n'.join(rtf_content)

            logger.info("PDF to RTF conversion completed")
            return ServiceResponse(
                status=200,
                message="PDF converted to RTF successfully",
                data=rtf_text.encode('utf-8'),
                format="rtf"
            )

        except Exception as e:
            logger.error("PDF to RTF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting PDF to RTF",
                error=str(e)
            )

    # TXT conversions
    async def convert_txt_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """Convert TXT to PDF."""
        try:
            if options is None:
                options = ConversionOptions()

            # Read text content
            text_content = file_buffer.decode('utf-8')
            
            # Create PDF
            pdf_buffer = io.BytesIO()
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            styles = getSampleStyleSheet()
            story = []

            for line in text_content.split('\n'):
                if line.strip():
                    p = Paragraph(self._clean_text(line), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info("TXT to PDF conversion completed")
            return ServiceResponse(
                status=200,
                message="TXT converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("TXT to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to PDF",
                error=str(e)
            )

    async def convert_txt_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert TXT to DOCX."""
        try:
            # Read text content
            text_content = file_buffer.decode('utf-8')
            
            # Create DOCX document
            doc = Document()
            for line in text_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()
            docx_buffer.close()

            logger.info("TXT to DOCX conversion completed")
            return ServiceResponse(
                status=200,
                message="TXT converted to DOCX successfully",
                data=docx_content,
                format="docx"
            )

        except Exception as e:
            logger.error("TXT to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to DOCX",
                error=str(e)
            )

    async def convert_txt_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert TXT to HTML."""
        try:
            # Read text content
            text_content = file_buffer.decode('utf-8')
            
            # Create HTML
            html_content = ['<html><head><title>Converted Text</title></head><body>']
            for line in text_content.split('\n'):
                if line.strip():
                    html_content.append(f'<p>{line}</p>')
            html_content.append('</body></html>')
            
            html_text = '\n'.join(html_content)

            logger.info("TXT to HTML conversion completed")
            return ServiceResponse(
                status=200,
                message="TXT converted to HTML successfully",
                data=html_text.encode('utf-8'),
                format="html"
            )

        except Exception as e:
            logger.error("TXT to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to HTML",
                error=str(e)
            )

    async def convert_txt_to_rtf(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert TXT to RTF."""
        try:
            # Read text content
            text_content = file_buffer.decode('utf-8')
            
            # Create RTF
            rtf_content = ['{\\rtf1\\ansi\\deff0']
            for line in text_content.split('\n'):
                if line.strip():
                    rtf_content.append(line)
                    rtf_content.append('\\par')
            rtf_content.append('}')
            
            rtf_text = '\n'.join(rtf_content)

            logger.info("TXT to RTF conversion completed")
            return ServiceResponse(
                status=200,
                message="TXT converted to RTF successfully",
                data=rtf_text.encode('utf-8'),
                format="rtf"
            )

        except Exception as e:
            logger.error("TXT to RTF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to RTF",
                error=str(e)
            )

    async def convert_txt_to_md(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert TXT to MD."""
        try:
            # Read text content
            text_content = file_buffer.decode('utf-8')
            
            # Create basic markdown
            md_content = []
            for line in text_content.split('\n'):
                if line.strip():
                    md_content.append(line)
            
            md_text = '\n'.join(md_content)

            logger.info("TXT to MD conversion completed")
            return ServiceResponse(
                status=200,
                message="TXT converted to MD successfully",
                data=md_text.encode('utf-8'),
                format="md"
            )

        except Exception as e:
            logger.error("TXT to MD conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to MD",
                error=str(e)
            )

    # HTML conversions
    async def convert_html_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """Convert HTML to PDF."""
        try:
            if options is None:
                options = ConversionOptions()

            # Parse HTML
            html_content = file_buffer.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text()
            
            # Create PDF
            pdf_buffer = io.BytesIO()
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            styles = getSampleStyleSheet()
            story = []

            for line in text_content.split('\n'):
                if line.strip():
                    p = Paragraph(self._clean_text(line), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info("HTML to PDF conversion completed")
            return ServiceResponse(
                status=200,
                message="HTML converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("HTML to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting HTML to PDF",
                error=str(e)
            )

    async def convert_html_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert HTML to DOCX."""
        try:
            html_content = file_buffer.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create DOCX document
            doc = Document()
            
            # Convert HTML elements to DOCX
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                text = element.get_text().strip()
                if text:
                    if element.name.startswith('h'):
                        level = int(element.name[1])
                        doc.add_heading(text, level=level)
                    else:
                        doc.add_paragraph(text)

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()
            docx_buffer.close()

            logger.info("HTML to DOCX conversion completed")
            return ServiceResponse(
                status=200,
                message="HTML converted to DOCX successfully",
                data=docx_content,
                format="docx"
            )

        except Exception as e:
            logger.error("HTML to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting HTML to DOCX",
                error=str(e)
            )

    async def convert_html_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert HTML to TXT."""
        try:
            html_content = file_buffer.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()

            logger.info("HTML to TXT conversion completed")
            return ServiceResponse(
                status=200,
                message="HTML converted to TXT successfully",
                data=text_content.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("HTML to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting HTML to TXT",
                error=str(e)
            )

    async def convert_html_to_md(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert HTML to Markdown."""
        try:
            html_content = file_buffer.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Convert to markdown-like format
            md_content = []
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                text = element.get_text().strip()
                if text:
                    if element.name.startswith('h'):
                        level = int(element.name[1])
                        md_content.append('#' * level + ' ' + text)
                    else:
                        md_content.append(text)
            
            md_text = '\n\n'.join(md_content)

            logger.info("HTML to MD conversion completed")
            return ServiceResponse(
                status=200,
                message="HTML converted to MD successfully",
                data=md_text.encode('utf-8'),
                format="md"
            )

        except Exception as e:
            logger.error("HTML to MD conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting HTML to MD",
                error=str(e)
            )

    # RTF conversions
    async def convert_rtf_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """Convert RTF to PDF."""
        try:
            if options is None:
                options = ConversionOptions()

            # Read RTF content
            rtf_content = file_buffer.decode('utf-8')
            
            # Strip RTF formatting (basic implementation)
            text_content = rtf_content.replace('\\par', '\n')
            text_content = text_content.replace('\\b', '')
            text_content = text_content.replace('\\i', '')
            text_content = text_content.replace('\\u', '')
            
            # Create PDF
            pdf_buffer = io.BytesIO()
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            styles = getSampleStyleSheet()
            story = []

            for line in text_content.split('\n'):
                if line.strip():
                    p = Paragraph(self._clean_text(line), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info("RTF to PDF conversion completed")
            return ServiceResponse(
                status=200,
                message="RTF converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("RTF to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting RTF to PDF",
                error=str(e)
            )

    async def convert_rtf_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert RTF to DOCX."""
        try:
            # Read RTF content
            rtf_content = file_buffer.decode('utf-8')
            
            # Strip RTF formatting (basic implementation)
            text_content = rtf_content.replace('\\par', '\n')
            text_content = text_content.replace('\\b', '')
            text_content = text_content.replace('\\i', '')
            text_content = text_content.replace('\\u', '')
            
            # Create DOCX document
            doc = Document()
            for line in text_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()
            docx_buffer.close()

            logger.info("RTF to DOCX conversion completed")
            return ServiceResponse(
                status=200,
                message="RTF converted to DOCX successfully",
                data=docx_content,
                format="docx"
            )

        except Exception as e:
            logger.error("RTF to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting RTF to DOCX",
                error=str(e)
            )

    async def convert_rtf_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert RTF to TXT."""
        try:
            # Read RTF content
            rtf_content = file_buffer.decode('utf-8')
            
            # Strip RTF formatting (basic implementation)
            text_content = rtf_content.replace('\\par', '\n')
            text_content = text_content.replace('\\b', '')
            text_content = text_content.replace('\\i', '')
            text_content = text_content.replace('\\u', '')

            logger.info("RTF to TXT conversion completed")
            return ServiceResponse(
                status=200,
                message="RTF converted to TXT successfully",
                data=text_content.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("RTF to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting RTF to TXT",
                error=str(e)
            )

    async def convert_rtf_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert RTF to HTML."""
        try:
            # Read RTF content
            rtf_content = file_buffer.decode('utf-8')
            
            # Strip RTF formatting (basic implementation)
            text_content = rtf_content.replace('\\par', '\n')
            text_content = text_content.replace('\\b', '')
            text_content = text_content.replace('\\i', '')
            text_content = text_content.replace('\\u', '')
            
            # Create HTML
            html_content = ['<html><head><title>Converted RTF</title></head><body>']
            for line in text_content.split('\n'):
                if line.strip():
                    html_content.append(f'<p>{line}</p>')
            html_content.append('</body></html>')
            
            html_text = '\n'.join(html_content)

            logger.info("RTF to HTML conversion completed")
            return ServiceResponse(
                status=200,
                message="RTF converted to HTML successfully",
                data=html_text.encode('utf-8'),
                format="html"
            )

        except Exception as e:
            logger.error("RTF to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting RTF to HTML",
                error=str(e)
            )

    # ODT conversions
    async def convert_odt_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """Convert ODT to PDF."""
        try:
            # For now, return a placeholder response
            logger.warning("ODT to PDF conversion requires additional libraries")
            return ServiceResponse(
                status=501,
                message="ODT to PDF conversion requires additional libraries (odfpy)",
                error="ODT to PDF conversion not fully implemented"
            )

        except Exception as e:
            logger.error("ODT to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting ODT to PDF",
                error=str(e)
            )

    async def convert_odt_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert ODT to DOCX."""
        try:
            # For now, return a placeholder response
            logger.warning("ODT to DOCX conversion requires additional libraries")
            return ServiceResponse(
                status=501,
                message="ODT to DOCX conversion requires additional libraries (odfpy)",
                error="ODT to DOCX conversion not fully implemented"
            )

        except Exception as e:
            logger.error("ODT to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting ODT to DOCX",
                error=str(e)
            )

    async def convert_odt_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert ODT to TXT."""
        try:
            # For now, return a placeholder response
            logger.warning("ODT to TXT conversion requires additional libraries")
            return ServiceResponse(
                status=501,
                message="ODT to TXT conversion requires additional libraries (odfpy)",
                error="ODT to TXT conversion not fully implemented"
            )

        except Exception as e:
            logger.error("ODT to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting ODT to TXT",
                error=str(e)
            )

    async def convert_odt_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert ODT to HTML."""
        try:
            # For now, return a placeholder response
            logger.warning("ODT to HTML conversion requires additional libraries")
            return ServiceResponse(
                status=501,
                message="ODT to HTML conversion requires additional libraries (odfpy)",
                error="ODT to HTML conversion not fully implemented"
            )

        except Exception as e:
            logger.error("ODT to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting ODT to HTML",
                error=str(e)
            )

    # Markdown conversions
    async def convert_md_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """Convert Markdown to PDF."""
        try:
            if options is None:
                options = ConversionOptions()

            md_content = file_buffer.decode('utf-8')
            
            # Convert markdown to HTML first
            html = markdown.markdown(md_content)
            
            # Then convert HTML to PDF
            html_buffer = io.BytesIO(html.encode('utf-8'))
            return await self.convert_html_to_pdf(html_buffer.getvalue(), options)

        except Exception as e:
            logger.error("MD to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting MD to PDF",
                error=str(e)
            )

    async def convert_md_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert Markdown to DOCX."""
        try:
            md_content = file_buffer.decode('utf-8')
            
            # Convert markdown to HTML first
            html = markdown.markdown(md_content)
            
            # Then convert HTML to DOCX
            html_buffer = io.BytesIO(html.encode('utf-8'))
            return await self.convert_html_to_docx(html_buffer.getvalue())

        except Exception as e:
            logger.error("MD to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting MD to DOCX",
                error=str(e)
            )

    async def convert_md_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert Markdown to TXT."""
        try:
            md_content = file_buffer.decode('utf-8')
            
            # Convert markdown to HTML first
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text_content = soup.get_text()

            logger.info("MD to TXT conversion completed")
            return ServiceResponse(
                status=200,
                message="MD converted to TXT successfully",
                data=text_content.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("MD to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting MD to TXT",
                error=str(e)
            )

    async def convert_md_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """Convert Markdown to HTML."""
        try:
            md_content = file_buffer.decode('utf-8')
            html = markdown.markdown(md_content)

            logger.info("MD to HTML conversion completed")
            return ServiceResponse(
                status=200,
                message="MD converted to HTML successfully",
                data=html.encode('utf-8'),
                format="html"
            )

        except Exception as e:
            logger.error("MD to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting MD to HTML",
                error=str(e)
            )

    async def get_supported_conversions(self):
        """Get list of supported document conversions."""
        return {
            "supported_conversions": self.supported_conversions,
            "message": "List of supported document format conversions"
        }


# Global instance
document_converter_service = DocumentConverterService()
