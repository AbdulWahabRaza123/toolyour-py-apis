"""
Document conversion service.
Handles DOCX to PDF and other document conversions.
"""

import io
import os
import tempfile
from typing import Optional
from docx import Document
from reportlab.lib.pagesizes import A4, letter, legal
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import PyPDF2
import structlog

# Image processing (removed - using separate image service)

# Office formats
from openpyxl import Workbook, load_workbook
from pptx import Presentation
import xlrd
import xlwt

# eBook formats
import ebooklib
from ebooklib import epub

# Audio/Video processing
from pydub import AudioSegment
# from moviepy.editor import VideoFileClip  # Temporarily disabled

# Archive handling
import zipfile
import rarfile
import py7zr

# Additional formats
import markdown
from bs4 import BeautifulSoup
import json

from .types import ServiceResponse, FileInput, ConversionOptions

logger = structlog.get_logger(__name__)


class DocumentConverterService:
    """Service for converting documents between formats."""

    def __init__(self):
        self.supported_conversions = {
            # Document formats
            'docx': ['pdf', 'txt', 'html', 'rtf', 'odt'],
            'pdf': ['txt', 'docx', 'html', 'rtf'],
            'txt': ['pdf', 'docx', 'html', 'rtf', 'md'],
            'html': ['pdf', 'docx', 'txt', 'md'],
            'rtf': ['pdf', 'docx', 'txt', 'html'],
            'odt': ['pdf', 'docx', 'txt', 'html'],
            'md': ['pdf', 'docx', 'txt', 'html'],
            
            # Office formats
            'xlsx': ['pdf', 'csv', 'xls', 'html', 'json'],
            'xls': ['pdf', 'csv', 'xlsx', 'html', 'json'],
            'csv': ['xlsx', 'xls', 'pdf', 'html', 'json'],
            'pptx': ['pdf', 'html', 'txt', 'pptx'],
            'ppt': ['pdf', 'html', 'txt', 'pptx'],
            
            # eBook formats
            'epub': ['pdf', 'txt', 'html', 'mobi'],
            'mobi': ['pdf', 'txt', 'html', 'epub'],
            
            # Archive formats
            'zip': ['rar', '7z', 'tar'],
            'rar': ['zip', '7z', 'tar'],
            '7z': ['zip', 'rar', 'tar'],
            
            # Audio formats
            'mp3': ['wav', 'flac', 'aac', 'ogg'],
            'wav': ['mp3', 'flac', 'aac', 'ogg'],
            'flac': ['mp3', 'wav', 'aac', 'ogg'],
            'aac': ['mp3', 'wav', 'flac', 'ogg'],
            'ogg': ['mp3', 'wav', 'flac', 'aac'],
            
            # Video formats
            'mp4': ['avi', 'mov', 'mkv', 'webm'],
            'avi': ['mp4', 'mov', 'mkv', 'webm'],
            'mov': ['mp4', 'avi', 'mkv', 'webm'],
            'mkv': ['mp4', 'avi', 'mov', 'webm'],
            'webm': ['mp4', 'avi', 'mov', 'mkv'],
        }

    def can_convert(self, source_format: str, target_format: str) -> bool:
        """Check if conversion is supported."""
        source_format = source_format.lower().replace('.', '')
        target_format = target_format.lower().replace('.', '')
        
        return target_format in self.supported_conversions.get(source_format, [])

    def get_supported_formats(self, source_format: str) -> list:
        """Get supported target formats for a source format."""
        source_format = source_format.lower().replace('.', '')
        return self.supported_conversions.get(source_format, [])

    async def convert_docx_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """
        Convert DOCX file to PDF.
        
        Args:
            file_buffer: The DOCX file content as bytes
            options: Conversion options (page size, orientation, etc.)
            
        Returns:
            ServiceResponse with PDF file buffer or error
        """
        try:
            if options is None:
                options = ConversionOptions()

            # Load DOCX document
            docx_stream = io.BytesIO(file_buffer)
            doc = Document(docx_stream)

            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            
            # Set page size
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            # Get styles
            styles = getSampleStyleSheet()
            story = []

            # Convert DOCX paragraphs to PDF
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Determine paragraph style based on DOCX style
                    style = self._get_paragraph_style(paragraph, styles)
                    
                    # Add paragraph to story
                    p = Paragraph(self._clean_text(paragraph.text), style)
                    story.append(p)
                    story.append(Spacer(1, 6))

            # Add tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        p = Paragraph(self._clean_text(row_text), styles['Normal'])
                        story.append(p)
                        story.append(Spacer(1, 3))
                story.append(Spacer(1, 12))

            # Build PDF
            pdf_doc.build(story)
            
            # Get PDF content
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info(
                "DOCX to PDF conversion completed",
                output_size=len(pdf_content)
            )

            return ServiceResponse(
                status=200,
                message="DOCX converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("DOCX to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to PDF",
                error=str(e)
            )

    async def convert_docx_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert DOCX file to plain text.
        
        Args:
            file_buffer: The DOCX file content as bytes
            
        Returns:
            ServiceResponse with text content or error
        """
        try:
            # Load DOCX document
            docx_stream = io.BytesIO(file_buffer)
            doc = Document(docx_stream)

            # Extract all text
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text_content.append(row_text)

            full_text = "\n".join(text_content)

            logger.info("DOCX to TXT conversion completed")

            return ServiceResponse(
                status=200,
                message="DOCX converted to TXT successfully",
                data=full_text.encode('utf-8'),
                format="txt"
            )

        except Exception as e:
            logger.error("DOCX to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to TXT",
                error=str(e)
            )

    async def convert_txt_to_pdf(
        self,
        file_buffer: bytes,
        options: Optional[ConversionOptions] = None
    ) -> ServiceResponse:
        """
        Convert plain text file to PDF.
        
        Args:
            file_buffer: The text file content as bytes
            options: Conversion options
            
        Returns:
            ServiceResponse with PDF file buffer or error
        """
        try:
            if options is None:
                options = ConversionOptions()

            # Decode text
            text_content = file_buffer.decode('utf-8')

            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            
            # Set page size
            page_size = self._get_page_size(options.page_size)
            if options.orientation == "landscape":
                page_size = (page_size[1], page_size[0])

            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=options.margin * mm,
                leftMargin=options.margin * mm,
                topMargin=options.margin * mm,
                bottomMargin=options.margin * mm,
            )

            # Get styles
            styles = getSampleStyleSheet()
            story = []

            # Convert text lines to PDF paragraphs
            for line in text_content.split('\n'):
                if line.strip():
                    p = Paragraph(self._clean_text(line), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))

            # Build PDF
            pdf_doc.build(story)
            
            # Get PDF content
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            logger.info("TXT to PDF conversion completed")

            return ServiceResponse(
                status=200,
                message="TXT converted to PDF successfully",
                data=pdf_content,
                format="pdf"
            )

        except Exception as e:
            logger.error("TXT to PDF conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to PDF",
                error=str(e)
            )

    def _get_page_size(self, size_name: str):
        """Get page size tuple based on name."""
        sizes = {
            'A4': A4,
            'letter': letter,
            'legal': legal,
        }
        return sizes.get(size_name, A4)

    def _get_paragraph_style(self, paragraph, styles):
        """Get appropriate PDF style for a DOCX paragraph."""
        # Check if it's a heading
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name.replace('Heading ', '')
            if level == '1':
                return styles['Heading1']
            elif level == '2':
                return styles['Heading2']
            elif level == '3':
                return styles['Heading3']
        
        # Check alignment
        if hasattr(paragraph, 'alignment') and paragraph.alignment:
            if paragraph.alignment == 1:  # Center
                return styles['Normal']
            elif paragraph.alignment == 2:  # Right
                return styles['Normal']
        
        return styles['Normal']

    async def convert_docx_to_html(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert DOCX file to HTML.
        
        Args:
            file_buffer: The DOCX file content as bytes
            
        Returns:
            ServiceResponse with HTML content or error
        """
        try:
            # Load DOCX document
            docx_stream = io.BytesIO(file_buffer)
            doc = Document(docx_stream)

            # Convert to HTML
            html_content = []
            html_content.append("<!DOCTYPE html>")
            html_content.append("<html><head><meta charset='utf-8'></head><body>")

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        html_content.append(f"<h{level}>{self._escape_html(paragraph.text)}</h{level}>")
                    else:
                        html_content.append(f"<p>{self._escape_html(paragraph.text)}</p>")

            # Add table content
            for table in doc.tables:
                html_content.append("<table border='1'>")
                for row in table.rows:
                    html_content.append("<tr>")
                    for cell in row.cells:
                        html_content.append(f"<td>{self._escape_html(cell.text)}</td>")
                    html_content.append("</tr>")
                html_content.append("</table>")

            html_content.append("</body></html>")
            full_html = "\n".join(html_content)

            logger.info("DOCX to HTML conversion completed")

            return ServiceResponse(
                status=200,
                message="DOCX converted to HTML successfully",
                data=full_html.encode('utf-8'),
                format="html"
            )

        except Exception as e:
            logger.error("DOCX to HTML conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting DOCX to HTML",
                error=str(e)
            )

    async def convert_pdf_to_txt(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert PDF file to plain text.
        
        Args:
            file_buffer: The PDF file content as bytes
            
        Returns:
            ServiceResponse with text content or error
        """
        try:
            # Create PDF reader from bytes
            pdf_stream = io.BytesIO(file_buffer)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            # Extract text from all pages
            text_content = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)
            
            # Join all text
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                logger.warning("No text content found in PDF")
                return ServiceResponse(
                    status=400,
                    message="No text content found in PDF file",
                    error="PDF appears to be image-based or contains no extractable text"
                )

            logger.info("PDF to TXT conversion completed", pages=len(pdf_reader.pages))

            return ServiceResponse(
                status=200,
                message="PDF converted to TXT successfully",
                data=full_text.encode('utf-8'),
                format="txt"
            )

        except PyPDF2.errors.PdfReadError as e:
            logger.error("Invalid PDF file", error=str(e))
            return ServiceResponse(
                status=400,
                message="Invalid PDF file format",
                error=f"PDF parsing error: {str(e)}"
            )
        except Exception as e:
            logger.error("PDF to TXT conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting PDF to TXT",
                error=str(e)
            )

    async def convert_txt_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert text file to DOCX.
        
        Args:
            file_buffer: The text file content as bytes
            
        Returns:
            ServiceResponse with DOCX file buffer or error
        """
        try:
            # Decode text
            text_content = file_buffer.decode('utf-8')

            # Create new DOCX document
            doc = Document()

            # Split text into paragraphs and add to document
            paragraphs = text_content.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()
            docx_buffer.close()

            logger.info("TXT to DOCX conversion completed")

            return ServiceResponse(
                status=200,
                message="TXT converted to DOCX successfully",
                data=docx_content,
                format="docx"
            )

        except Exception as e:
            logger.error("TXT to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting TXT to DOCX",
                error=str(e)
            )

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&#x27;')
        return text

    async def convert_pdf_to_docx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert PDF file to DOCX.
        
        Args:
            file_buffer: The PDF file content as bytes
            
        Returns:
            ServiceResponse with DOCX file buffer or error
        """
        try:
            # Create PDF reader from bytes
            pdf_stream = io.BytesIO(file_buffer)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            # Extract text from all pages
            text_content = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)
            
            # Join all text
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                logger.warning("No text content found in PDF")
                return ServiceResponse(
                    status=400,
                    message="No text content found in PDF file",
                    error="PDF appears to be image-based or contains no extractable text"
                )

            # Create new DOCX document
            doc = Document()

            # Split text into paragraphs and add to document
            paragraphs = full_text.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()
            docx_buffer.close()

            logger.info("PDF to DOCX conversion completed", pages=len(pdf_reader.pages))

            return ServiceResponse(
                status=200,
                message="PDF converted to DOCX successfully",
                data=docx_content,
                format="docx"
            )

        except PyPDF2.errors.PdfReadError as e:
            logger.error("Invalid PDF file", error=str(e))
            return ServiceResponse(
                status=400,
                message="Invalid PDF file format",
                error=f"PDF parsing error: {str(e)}"
            )
        except Exception as e:
            logger.error("PDF to DOCX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting PDF to DOCX",
                error=str(e)
            )


    async def convert_xlsx_to_csv(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert XLSX to CSV.
        
        Args:
            file_buffer: The XLSX file content as bytes
            
        Returns:
            ServiceResponse with CSV content or error
        """
        try:
            # Load workbook
            workbook = load_workbook(io.BytesIO(file_buffer))
            worksheet = workbook.active
            
            # Convert to CSV
            csv_content = []
            for row in worksheet.iter_rows(values_only=True):
                csv_content.append(','.join(str(cell) if cell is not None else '' for cell in row))
            
            csv_text = '\n'.join(csv_content)
            
            logger.info("XLSX to CSV conversion completed")
            
            return ServiceResponse(
                status=200,
                message="XLSX converted to CSV successfully",
                data=csv_text.encode('utf-8'),
                format="csv"
            )
            
        except Exception as e:
            logger.error("XLSX to CSV conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting XLSX to CSV",
                error=str(e)
            )

    async def convert_csv_to_xlsx(
        self,
        file_buffer: bytes
    ) -> ServiceResponse:
        """
        Convert CSV to XLSX.
        
        Args:
            file_buffer: The CSV file content as bytes
            
        Returns:
            ServiceResponse with XLSX file buffer or error
        """
        try:
            # Read CSV content
            csv_text = file_buffer.decode('utf-8')
            lines = csv_text.strip().split('\n')
            
            # Create workbook
            workbook = Workbook()
            worksheet = workbook.active
            
            # Add data to worksheet
            for row_idx, line in enumerate(lines, 1):
                cells = line.split(',')
                for col_idx, cell_value in enumerate(cells, 1):
                    worksheet.cell(row=row_idx, column=col_idx, value=cell_value)
            
            # Save to bytes
            xlsx_buffer = io.BytesIO()
            workbook.save(xlsx_buffer)
            xlsx_content = xlsx_buffer.getvalue()
            xlsx_buffer.close()
            
            logger.info("CSV to XLSX conversion completed")
            
            return ServiceResponse(
                status=200,
                message="CSV converted to XLSX successfully",
                data=xlsx_content,
                format="xlsx"
            )
            
        except Exception as e:
            logger.error("CSV to XLSX conversion failed", error=str(e))
            return ServiceResponse(
                status=500,
                message="Error converting CSV to XLSX",
                error=str(e)
            )

    async def convert_audio_to_audio(
        self,
        file_buffer: bytes,
        target_format: str
    ) -> ServiceResponse:
        """
        Convert audio from one format to another.
        
        Args:
            file_buffer: The audio file content as bytes
            target_format: Target audio format
            
        Returns:
            ServiceResponse with converted audio or error
        """
        try:
            # Load audio
            audio = AudioSegment.from_file(io.BytesIO(file_buffer))
            
            # Export to target format
            output_buffer = io.BytesIO()
            audio.export(output_buffer, format=target_format.lower())
            converted_data = output_buffer.getvalue()
            output_buffer.close()
            
            logger.info(f"Audio conversion completed to {target_format}")
            
            return ServiceResponse(
                status=200,
                message=f"Audio converted to {target_format.upper()} successfully",
                data=converted_data,
                format=target_format.lower()
            )
            
        except Exception as e:
            logger.error(f"Audio conversion failed to {target_format}", error=str(e))
            return ServiceResponse(
                status=500,
                message=f"Error converting audio to {target_format}",
                error=str(e)
            )

            def _clean_text(self, text: str) -> str:
                """Clean text for PDF generation (escape XML special characters)."""
                text = text.replace('&', '&amp;')
                text = text.replace('<', '&lt;')
                text = text.replace('>', '&gt;')
                return text

            # HTML conversions
            async def convert_html_to_pdf(
                self,
                file_buffer: bytes,
                options: Optional[ConversionOptions] = None
            ) -> ServiceResponse:
                """Convert HTML to PDF."""
                try:
                    if options is None:
                        options = ConversionOptions()

                    # Parse HTML
                    html_content = file_buffer.decode('utf-8')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # Extract text content
                    text_content = soup.get_text()
                    
                    # Create PDF
                    pdf_buffer = io.BytesIO()
                    page_size = self._get_page_size(options.page_size)
                    if options.orientation == "landscape":
                        page_size = (page_size[1], page_size[0])

                    pdf_doc = SimpleDocTemplate(
                        pdf_buffer,
                        pagesize=page_size,
                        rightMargin=options.margin * mm,
                        leftMargin=options.margin * mm,
                        topMargin=options.margin * mm,
                        bottomMargin=options.margin * mm,
                    )

                    styles = getSampleStyleSheet()
                    story = []

                    for line in text_content.split('\n'):
                        if line.strip():
                            p = Paragraph(self._clean_text(line), styles['Normal'])
                            story.append(p)
                            story.append(Spacer(1, 6))

                    pdf_doc.build(story)
                    pdf_content = pdf_buffer.getvalue()
                    pdf_buffer.close()

                    logger.info("HTML to PDF conversion completed")
                    return ServiceResponse(
                        status=200,
                        message="HTML converted to PDF successfully",
                        data=pdf_content,
                        format="pdf"
                    )

                except Exception as e:
                    logger.error("HTML to PDF conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting HTML to PDF",
                        error=str(e)
                    )

            async def convert_html_to_docx(
                self,
                file_buffer: bytes
            ) -> ServiceResponse:
                """Convert HTML to DOCX."""
                try:
                    html_content = file_buffer.decode('utf-8')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # Create DOCX document
                    doc = Document()
                    
                    # Convert HTML elements to DOCX
                    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                        text = element.get_text().strip()
                        if text:
                            if element.name.startswith('h'):
                                level = int(element.name[1])
                                heading_style = f'Heading {level}'
                                doc.add_heading(text, level=level)
                            else:
                                doc.add_paragraph(text)

                    # Save to bytes
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_content = docx_buffer.getvalue()
                    docx_buffer.close()

                    logger.info("HTML to DOCX conversion completed")
                    return ServiceResponse(
                        status=200,
                        message="HTML converted to DOCX successfully",
                        data=docx_content,
                        format="docx"
                    )

                except Exception as e:
                    logger.error("HTML to DOCX conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting HTML to DOCX",
                        error=str(e)
                    )

            async def convert_html_to_txt(
                self,
                file_buffer: bytes
            ) -> ServiceResponse:
                """Convert HTML to TXT."""
                try:
                    html_content = file_buffer.decode('utf-8')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    text_content = soup.get_text()

                    logger.info("HTML to TXT conversion completed")
                    return ServiceResponse(
                        status=200,
                        message="HTML converted to TXT successfully",
                        data=text_content.encode('utf-8'),
                        format="txt"
                    )

                except Exception as e:
                    logger.error("HTML to TXT conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting HTML to TXT",
                        error=str(e)
                    )

            async def convert_html_to_md(
                self,
                file_buffer: bytes
            ) -> ServiceResponse:
                """Convert HTML to Markdown."""
                try:
                    html_content = file_buffer.decode('utf-8')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # Convert to markdown-like format
                    md_content = []
                    
                    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                        text = element.get_text().strip()
                        if text:
                            if element.name.startswith('h'):
                                level = int(element.name[1])
                                md_content.append('#' * level + ' ' + text)
                            else:
                                md_content.append(text)
                    
                    md_text = '\n\n'.join(md_content)

                    logger.info("HTML to MD conversion completed")
                    return ServiceResponse(
                        status=200,
                        message="HTML converted to MD successfully",
                        data=md_text.encode('utf-8'),
                        format="md"
                    )

                except Exception as e:
                    logger.error("HTML to MD conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting HTML to MD",
                        error=str(e)
                    )

            # Markdown conversions
            async def convert_md_to_pdf(
                self,
                file_buffer: bytes,
                options: Optional[ConversionOptions] = None
            ) -> ServiceResponse:
                """Convert Markdown to PDF."""
                try:
                    if options is None:
                        options = ConversionOptions()

                    md_content = file_buffer.decode('utf-8')
                    
                    # Convert markdown to HTML first
                    html = markdown.markdown(md_content)
                    
                    # Then convert HTML to PDF
                    html_buffer = io.BytesIO(html.encode('utf-8'))
                    return await self.convert_html_to_pdf(html_buffer.getvalue(), options)

                except Exception as e:
                    logger.error("MD to PDF conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting MD to PDF",
                        error=str(e)
                    )

            async def convert_md_to_docx(
                self,
                file_buffer: bytes
            ) -> ServiceResponse:
                """Convert Markdown to DOCX."""
                try:
                    md_content = file_buffer.decode('utf-8')
                    
                    # Convert markdown to HTML first
                    html = markdown.markdown(md_content)
                    
                    # Then convert HTML to DOCX
                    html_buffer = io.BytesIO(html.encode('utf-8'))
                    return await self.convert_html_to_docx(html_buffer.getvalue())

                except Exception as e:
                    logger.error("MD to DOCX conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting MD to DOCX",
                        error=str(e)
                    )

            async def convert_md_to_txt(
                self,
                file_buffer: bytes
            ) -> ServiceResponse:
                """Convert Markdown to TXT."""
                try:
                    md_content = file_buffer.decode('utf-8')
                    
                    # Convert markdown to HTML first
                    html = markdown.markdown(md_content)
                    soup = BeautifulSoup(html, 'html.parser')
                    text_content = soup.get_text()

                    logger.info("MD to TXT conversion completed")
                    return ServiceResponse(
                        status=200,
                        message="MD converted to TXT successfully",
                        data=text_content.encode('utf-8'),
                        format="txt"
                    )

                except Exception as e:
                    logger.error("MD to TXT conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting MD to TXT",
                        error=str(e)
                    )

            async def convert_md_to_html(
                self,
                file_buffer: bytes
            ) -> ServiceResponse:
                """Convert Markdown to HTML."""
                try:
                    md_content = file_buffer.decode('utf-8')
                    html = markdown.markdown(md_content)

                    logger.info("MD to HTML conversion completed")
                    return ServiceResponse(
                        status=200,
                        message="MD converted to HTML successfully",
                        data=html.encode('utf-8'),
                        format="html"
                    )

                except Exception as e:
                    logger.error("MD to HTML conversion failed", error=str(e))
                    return ServiceResponse(
                        status=500,
                        message="Error converting MD to HTML",
                        error=str(e)
                    )


# Global instance
document_converter_service = DocumentConverterService()

